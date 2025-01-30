from flask import Flask, jsonify, request, send_file, send_from_directory
from flask_cors import CORS
from supabase import create_client
from dotenv import load_dotenv
import os
from datetime import datetime
from models.models import Company, Employee, Client, Quotation, Item
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO
from werkzeug.utils import secure_filename
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import base64
import requests
from collections import OrderedDict
from docxtpl import DocxTemplate

# Load environment variables
load_dotenv()

app = Flask(__name__)
# Enable CORS for all routes during development
CORS(app, 
     resources={r"/api/*": {
         "origins": ["http://localhost:3000"],
         "supports_credentials": True,
         "allow_headers": ["Content-Type", "Authorization"],
         "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"]
     }})

# Configure Flask app
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY')
app.config['DEBUG'] = True

# Configure upload folder
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Initialize Supabase client
try:
    supabase_url = os.getenv('SUPABASE_URL')
    supabase_key = os.getenv('SUPABASE_KEY')
    
    if not supabase_url or not supabase_key:
        raise ValueError("Supabase URL or Key not found in environment variables")
    
    supabase = create_client(supabase_url, supabase_key)
    
    # Test the connection
    test_response = supabase.table('companies').select('*').limit(1).execute()
    print("Successfully connected to Supabase!")
except ValueError as e:
    print(f"Configuration Error: {e}")
    raise
except Exception as e:
    print(f"Error connecting to Supabase: {str(e)}")
    raise

# Health check route
@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy", "message": "API is running"}), 200

# Company routes
@app.route('/api/companies', methods=['GET'])
def get_companies():
    try:
        companies = supabase.table('companies').select('*').execute()
        return jsonify({
            "success": True,
            "data": [Company.from_db(company) for company in companies.data]
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/companies', methods=['POST'])
def create_company():
    try:
        data = request.json
        company = supabase.table('companies').insert(data).execute()
        return jsonify({
            "success": True,
            "data": Company.from_db(company.data[0])
        }), 201
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/companies/<int:company_id>', methods=['PUT'])
def update_company(company_id):
    try:
        data = request.json
        company = supabase.table('companies').update(data).eq('id', company_id).execute()
        if not company.data:
            return jsonify({"success": False, "error": "Company not found"}), 404
        return jsonify({
            "success": True,
            "data": Company.from_db(company.data[0])
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/companies/<int:company_id>', methods=['DELETE'])
def delete_company(company_id):
    try:
        # First delete associated quotations
        quotations = supabase.table('quotations').delete().eq('company_id', company_id).execute()
        
        # Then delete associated clients
        clients = supabase.table('clients').delete().eq('company_id', company_id).execute()
        
        # Finally delete the company
        company = supabase.table('companies').delete().eq('id', company_id).execute()
        
        if not company.data:
            return jsonify({"success": False, "error": "Company not found"}), 404
            
        return jsonify({
            "success": True,
            "message": "Company and all associated data deleted successfully"
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/companies/<int:company_id>/seal', methods=['POST'])
def upload_company_seal(company_id):
    try:
        if 'seal_image' not in request.files:
            return jsonify({"success": False, "error": "No file provided"}), 400
            
        file = request.files['seal_image']
        if file.filename == '':
            return jsonify({"success": False, "error": "No file selected"}), 400
            
        if file:
            # Secure the filename and create unique name
            filename = secure_filename(file.filename)
            unique_filename = f"{company_id}_{int(datetime.now().timestamp())}_{filename}"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            
            # Save the file
            file.save(filepath)
            
            # Update company with seal image URL
            seal_url = f"/uploads/{unique_filename}"  # URL path to access the image
            company = supabase.table('companies').update({
                'seal_image_url': seal_url
            }).eq('id', company_id).execute()
            
            if not company.data:
                return jsonify({"success": False, "error": "Company not found"}), 404
                
            return jsonify({
                "success": True,
                "data": Company.from_db(company.data[0])
            })
            
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# Serve uploaded files
@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename))

# Client routes
@app.route('/api/clients', methods=['GET'])
def get_clients():
    try:
        company_id = request.args.get('company_id')
        query = supabase.table('clients').select('*')
        if company_id:
            query = query.eq('company_id', company_id)
        clients = query.execute()
        return jsonify({
            "success": True,
            "data": [Client.from_db(client) for client in clients.data]
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/clients', methods=['POST'])
def create_client():
    try:
        data = request.json
        client = supabase.table('clients').insert(data).execute()
        return jsonify({
            "success": True,
            "data": Client.from_db(client.data[0])
        }), 201
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/clients/<int:client_id>', methods=['PUT'])
def update_client(client_id):
    try:
        data = request.json
        client = supabase.table('clients').update(data).eq('id', client_id).execute()
        if not client.data:
            return jsonify({"success": False, "error": "Client not found"}), 404
        return jsonify({
            "success": True,
            "data": Client.from_db(client.data[0])
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# Employee routes
@app.route('/api/employees', methods=['GET'])
def get_employees():
    try:
        employees = supabase.table('employees').select('*').execute()
        return jsonify({
            "success": True,
            "data": [Employee.from_db(employee) for employee in employees.data]
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/employees', methods=['POST'])
def create_employee():
    try:
        data = request.json
        employee = supabase.table('employees').insert(data).execute()
        return jsonify({
            "success": True,
            "data": Employee.from_db(employee.data[0])
        }), 201
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/employees/<int:employee_id>', methods=['PUT'])
def update_employee(employee_id):
    try:
        data = request.json
        employee = supabase.table('employees').update(data).eq('id', employee_id).execute()
        if not employee.data:
            return jsonify({"success": False, "error": "Employee not found"}), 404
        return jsonify({
            "success": True,
            "data": Employee.from_db(employee.data[0])
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/employees/<int:employee_id>', methods=['DELETE'])
def delete_employee(employee_id):
    try:
        # Check if employee has any quotations
        quotations = supabase.table('quotations').select('id').eq('employee_id', employee_id).execute()
        if quotations.data:
            return jsonify({
                "success": False,
                "error": "Cannot delete employee with existing quotations"
            }), 400

        employee = supabase.table('employees').delete().eq('id', employee_id).execute()
        if not employee.data:
            return jsonify({"success": False, "error": "Employee not found"}), 404
            
        return jsonify({
            "success": True,
            "message": "Employee deleted successfully"
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# Quotation routes
@app.route('/api/quotations', methods=['GET'])
def get_quotations():
    try:
        # Get quotations with company and client information only
        quotations = supabase.table('quotations') \
            .select('id, ref_number, date, total, companies(name), clients(name)') \
            .execute()

        # Process the results
        processed_quotations = []
        for quotation in quotations.data:
            company_name = quotation['companies']['name'] if quotation.get('companies') else None
            client_name = quotation['clients']['name'] if quotation.get('clients') else None
            
            processed_quotation = {
                'id': quotation['id'],
                'ref_number': quotation['ref_number'],
                'company': company_name,
                'client': client_name,
                'date': quotation['date'],
                'total': float(quotation['total']) if quotation.get('total') else 0
            }
            
            processed_quotations.append(processed_quotation)

        return jsonify({
            "success": True,
            "data": processed_quotations
        })
    except Exception as e:
        print(f"Error in get_quotations: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/quotations/<int:quotation_id>', methods=['DELETE'])
def delete_quotation(quotation_id):
    try:
        # Delete the quotation
        result = supabase.table('quotations').delete().eq('id', quotation_id).execute()
        
        if not result.data:
            return jsonify({
                'success': False,
                'error': 'Quotation not found'
            }), 404
            
        return jsonify({
            'success': True,
            'message': 'Quotation deleted successfully'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/quotations', methods=['POST'])
def create_quotation():
    try:
        data = request.json
        print("Received data:", data)  # Debug log
        
        # Fetch company data
        company_response = supabase.table('companies').select('*').eq('id', data['company_id']).execute()
        if not company_response.data:
            return jsonify({"success": False, "error": "Company not found"}), 404
            
        company = Company.from_db(company_response.data[0])
        
        # Generate reference number
        new_number = company['last_quote_number'] + 1
        ref_number = company['ref_format'].format(
            YYYY=datetime.now().year,
            NUM=str(new_number).zfill(4)
        )
        
        # Prepare quotation data
        quotation_data = {
            'company_id': data['company_id'],
            'client_id': data['client_id'],
            'ref_number': ref_number,
            'date': datetime.utcnow().isoformat(),
            'items': data['items'],
            'total': data['total']
        }
        
        # Create quotation
        quotation_response = supabase.table('quotations').insert(quotation_data).execute()
        
        if not quotation_response.data:
            raise Exception("Failed to create quotation")
            
        # Update company's last quote number
        company_update = supabase.table('companies').update({
            'last_quote_number': new_number
        }).eq('id', data['company_id']).execute()
        
        if not company_update.data:
            print("Warning: Failed to update company's last quote number")
            
        return jsonify({
            "success": True,
            "data": Quotation.from_db(quotation_response.data[0])
        }), 201
            
    except Exception as e:
        print("Error creating quotation:", str(e))  # Debug log
        return jsonify({"success": False, "error": str(e)}), 500

# Document generation route
@app.route('/api/generate-quote/<int:quotation_id>', methods=['GET'])
def generate_quote(quotation_id):
    try:
        # Fetch quotation data
        quotation = supabase.table('quotations').select('*').eq('id', quotation_id).execute()
        if not quotation.data:
            return jsonify({"success": False, "error": "Quotation not found"}), 404
        
        quotation_data = quotation.data[0]
        
        # Fetch company data
        company = supabase.table('companies').select('*').eq('id', quotation_data['company_id']).execute()
        if not company.data:
            return jsonify({"success": False, "error": "Company not found"}), 404
            
        company_data = company.data[0]
        
        # Create a new document
        doc = Document()
        
        # Add header table
        header_table = doc.add_table(rows=4, cols=1)
        header_table.style = 'Table Grid'
        
        # Company Name Cell
        company_cell = header_table.rows[0].cells[0]
        company_name = company_cell.paragraphs[0]
        company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_name.add_run(company_data['name'].upper())
        company_name.runs[0].font.size = Pt(16)
        company_name.runs[0].font.bold = True
        company_name.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Address Cell
        address_cell = header_table.rows[1].cells[0]
        address = address_cell.paragraphs[0]
        address.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address.add_run(company_data['address'].upper())
        address.runs[0].font.size = Pt(11)
        address.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Email and Phone Cell
        email_cell = header_table.rows[2].cells[0]
        email = email_cell.paragraphs[0]
        email.alignment = WD_ALIGN_PARAGRAPH.CENTER
        email.add_run(f"Email:- {company_data['email']} {company_data['phone']}")
        email.runs[0].font.size = Pt(11)
        email.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Tax Info Cell
        tax_cell = header_table.rows[3].cells[0]
        tax_info = tax_cell.paragraphs[0]
        tax_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tax_info.add_run(f"PAN NO.: {company_data['pan']} | GST NO.: {company_data['gst']}")
        tax_info.runs[0].font.size = Pt(11)
        tax_info.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Set background color for all cells and remove borders
        for row in header_table.rows:
            for cell in row.cells:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4F8C"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
                # Remove cell borders
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>' +
                                    '<w:top w:val="nil"/>' +
                                    '<w:left w:val="nil"/>' +
                                    '<w:bottom w:val="nil"/>' +
                                    '<w:right w:val="nil"/>' +
                                    '</w:tcBorders>')
                tcPr.append(tcBorders)

        # Add QUOTATION/PERFORMA INVOICE title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run('\nQUOTATION/PERFORMA INVOICE\n')
        title.runs[0].font.bold = True
        title.runs[0].font.size = Pt(12)

        # Add reference number and date
        ref_date = doc.add_table(rows=1, cols=2)
        ref_date.autofit = True
        ref_cell = ref_date.cell(0, 0)
        ref_cell.text = f"Ref No: {quotation_data['ref_number']}"
        date_cell = ref_date.cell(0, 1)
        date_cell.text = f"Date: {quotation_data['date'][:10]}"
        date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add spacing after header
        doc.add_paragraph()

        # Set default font size for the document
        style = doc.styles['Normal']
        style.font.size = Pt(9)

        # Add client details
        to_table = doc.add_table(rows=2, cols=1)
        to_table.style = 'Table Grid'
        
        # To Cell with blue background
        to_cell = to_table.rows[0].cells[0]
        to_paragraph = to_cell.paragraphs[0]
        to_paragraph.add_run('To')
        to_paragraph.runs[0].font.bold = True
        to_paragraph.runs[0].font.size = Pt(9)
        to_paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Set blue background for To cell
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4F8C"/>')
        to_cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Client details cell
        client = data.get('client', {})
        details_cell = to_table.rows[1].cells[0]
        details_paragraph = details_cell.paragraphs[0]
        details_paragraph.add_run(f"{client.get('business_name', '')}\n")
        details_paragraph.add_run(f"{client.get('address', '')}\n")
        details_paragraph.add_run(f"Kind Attn: {client.get('name', '')} | Tel: {client.get('phone', '')} | Email: {client.get('email', '')}")
        
        # Remove borders from both cells
        for row in to_table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>' +
                                    '<w:top w:val="nil"/>' +
                                    '<w:left w:val="nil"/>' +
                                    '<w:bottom w:val="nil"/>' +
                                    '<w:right w:val="nil"/>' +
                                    '</w:tcBorders>')
                tcPr.append(tcBorders)
        
        # Add spacing after client details
        doc.add_paragraph()
        
        # Add greeting text
        greeting = doc.add_paragraph()
        greeting.add_run("Dear Sir/Madam,\n")
        greeting.add_run("Thank you for your enquiry. We are pleased to quote our best prices as under:\n")
        
        # Add spacing after greeting
        doc.add_paragraph()
        
        # Add items table
        doc.add_paragraph().add_run('Items:').bold = True
        table = doc.add_table(rows=1, cols=14)  # Changed to 14 columns
        table.style = 'Table Grid'
        
        # Set header row
        header_cells = table.rows[0].cells
        headers = ['S.No', 'Cat No.', 'Description', 'Pack Size', 'HSN Code', 'Qty', 'Unit Rate', 'Discounted Price', 'Expanded Price', 'GST %', 'GST', 'Total Value', 'Lead Time', 'Brand']
        for i, text in enumerate(headers):
            header_cells[i].text = text
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)  # Set header font size
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add item rows
        items = data.get('items', [])
        for idx, item in enumerate(items, 1):
            row = table.add_row().cells
            row[0].text = str(idx)  # S.No
            row[1].text = item.get('catalogue_id', '')  # Cat No.
            row[2].text = item.get('description', '')  # Description
            row[3].text = item.get('pack_size', '')  # Pack Size
            row[4].text = item.get('hsn', '')  # HSN Code
            row[5].text = str(item.get('quantity', ''))  # Qty
            row[6].text = f"₹{item.get('unit_rate', 0):.2f}"  # Unit Rate
            
            # Calculate discounted price
            unit_rate = float(item.get('unit_rate', 0))
            discount = float(item.get('discount_percentage', 0))
            discounted_price = unit_rate * (1 - discount/100)
            row[7].text = f"₹{discounted_price:.2f}"  # Discounted Price
            
            # Calculate expanded price (discounted price * quantity)
            quantity = float(item.get('quantity', 0))
            expanded_price = discounted_price * quantity
            row[8].text = f"₹{expanded_price:.2f}"  # Expanded Price
            
            row[9].text = f"{item.get('gst_percentage', 0)}%"  # GST %
            row[10].text = f"₹{item.get('gst_value', 0):.2f}"  # GST
            row[11].text = f"₹{item.get('total', 0):.2f}"  # Total Value
            row[12].text = item.get('lead_time', '')  # Lead Time
            row[13].text = item.get('brand', '')  # Changed from 'make' to 'brand'
            
            # Center align all cells
            for cell in row:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set optimized column widths for better fit
        for cell in table.columns[0].cells:  # S.No
            cell.width = Inches(0.2)
        for cell in table.columns[1].cells:  # Cat No.
            cell.width = Inches(0.6)
        for cell in table.columns[2].cells:  # Description
            cell.width = Inches(1.8)
        for cell in table.columns[3].cells:  # Pack Size
            cell.width = Inches(0.4)
        for cell in table.columns[4].cells:  # HSN Code
            cell.width = Inches(0.6)
        for cell in table.columns[5].cells:  # Qty
            cell.width = Inches(0.3)
        for cell in table.columns[6].cells:  # Unit Rate
            cell.width = Inches(0.6)
        for cell in table.columns[7].cells:  # Discounted Price
            cell.width = Inches(0.6)
        for cell in table.columns[8].cells:  # Expanded Price
            cell.width = Inches(0.6)
        for cell in table.columns[9].cells:  # GST %
            cell.width = Inches(0.4)
        for cell in table.columns[10].cells:  # GST
            cell.width = Inches(0.6)
        for cell in table.columns[11].cells:  # Total Value
            cell.width = Inches(0.6)
        for cell in table.columns[12].cells:  # Lead Time
            cell.width = Inches(0.6)
        for cell in table.columns[13].cells:  # Brand
            cell.width = Inches(0.5)
        
        # Add totals
        doc.add_paragraph()
        totals = doc.add_paragraph()
        totals.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        totals.add_run(f"Sub Total: ₹{data.get('total', 0):.2f}\n")
        totals.add_run(f"Total GST: ₹{data.get('totalGST', 0):.2f}\n")
        grand_total = totals.add_run(f"Grand Total: ₹{data.get('total', 0):.2f}")
        grand_total.bold = True
        
        # Add terms and conditions
        terms_table = doc.add_table(rows=1, cols=1)
        terms_table.style = 'Table Grid'
        
        # Terms & Conditions header cell with blue background
        header_cell = terms_table.rows[0].cells[0]
        header_paragraph = header_cell.paragraphs[0]
        header_run = header_paragraph.add_run('Terms & Conditions')
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(255, 255, 255)
        header_run.font.size = Pt(11)
        
        # Set blue background for header cell
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4F8C"/>')
        header_cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Remove borders from the header cell
        tcPr = header_cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>' +
                            '<w:top w:val="nil"/>' +
                            '<w:left w:val="nil"/>' +
                            '<w:bottom w:val="nil"/>' +
                            '<w:right w:val="nil"/>' +
                            '</w:tcBorders>')
        tcPr.append(tcBorders)
        
        # Add payment terms and other terms as bullet points
        terms_list = doc.add_paragraph()
        terms_list.style = doc.styles['Normal']
        terms_list.paragraph_format.space_before = Pt(6)
        
        # Add payment terms
        payment_term = terms_list.add_run(f"1) {data.get('paymentTerms', '')}\n")
        payment_term.font.size = Pt(8)
        
        # Add fixed terms with numbering
        for idx, term in enumerate(data.get('fixedTerms', []), 2):
            term_text = terms_list.add_run(f"{idx}) {term}\n")
            term_text.font.size = Pt(8)
        
        # Add spacing
        doc.add_paragraph('\n')
        
        # Add Bank Details section
        bank_table = doc.add_table(rows=1, cols=1)
        bank_table.style = 'Table Grid'
        
        # Bank Details header cell with blue background
        bank_header_cell = bank_table.rows[0].cells[0]
        bank_header_paragraph = bank_header_cell.paragraphs[0]
        bank_header_run = bank_header_paragraph.add_run('Bank Details')
        bank_header_run.font.bold = True
        bank_header_run.font.color.rgb = RGBColor(255, 255, 255)
        bank_header_run.font.size = Pt(11)
        
        # Set blue background for header cell
        bank_shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4F8C"/>')
        bank_header_cell._tc.get_or_add_tcPr().append(bank_shading_elm)
        
        # Remove borders from the header cell
        bank_tcPr = bank_header_cell._tc.get_or_add_tcPr()
        bank_tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>' +
                            '<w:top w:val="nil"/>' +
                            '<w:left w:val="nil"/>' +
                            '<w:bottom w:val="nil"/>' +
                            '<w:right w:val="nil"/>' +
                            '</w:tcBorders>')
        bank_tcPr.append(bank_tcBorders)
        
        # Add bank details content
        bank_details = doc.add_paragraph()
        bank_details.style = doc.styles['Normal']
        bank_details.paragraph_format.space_before = Pt(6)
        bank_details_text = bank_details.add_run(f"HDFC BANK LTD. Account No: {data.get('company', {}).get('account_number', '')} ; NEFT/RTGS IFCS : {data.get('company', {}).get('ifsc_code', '')} Branch code:{data.get('company', {}).get('branch_code', '')} ; Micro code : {data.get('company', {}).get('micro_code', '')} ;Account type: Current account")
        bank_details_text.font.size = Pt(8)
        
        # Add spacing before quotation created by section
        doc.add_paragraph('\n')
        
        # Add Quotation Created By section
        created_by_table = doc.add_table(rows=1, cols=1)
        created_by_table.style = 'Table Grid'
        
        # Created By header cell with blue background
        created_by_header_cell = created_by_table.rows[0].cells[0]
        created_by_header_paragraph = created_by_header_cell.paragraphs[0]
        created_by_header_run = created_by_header_paragraph.add_run('Quotation Created By')
        created_by_header_run.font.bold = True
        created_by_header_run.font.color.rgb = RGBColor(255, 255, 255)
        created_by_header_run.font.size = Pt(11)
        
        # Set blue background for header cell
        created_by_shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4F8C"/>')
        created_by_header_cell._tc.get_or_add_tcPr().append(created_by_shading_elm)
        
        # Remove borders from the header cell
        created_by_tcPr = created_by_header_cell._tc.get_or_add_tcPr()
        created_by_tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>' +
                            '<w:top w:val="nil"/>' +
                            '<w:left w:val="nil"/>' +
                            '<w:bottom w:val="nil"/>' +
                            '<w:right w:val="nil"/>' +
                            '</w:tcBorders>')
        created_by_tcPr.append(created_by_tcBorders)
        
        # Add employee details content
        employee_details = doc.add_paragraph()
        employee_details.style = doc.styles['Normal']
        employee_details.paragraph_format.space_before = Pt(6)
        employee_name = data.get('employee', {}).get('name', '')
        employee_phone = data.get('employee', {}).get('phone_number', '')  # Changed from mobile to phone_number
        employee_email = data.get('employee', {}).get('email', '')
        
        employee_details.add_run(f"{employee_name}\n").font.size = Pt(8)
        employee_details.add_run(f"Mobile: {employee_phone}\n").font.size = Pt(8)  # Using the new employee_phone variable
        email_run = employee_details.add_run(f"Email: ")
        email_run.font.size = Pt(8)
        email_link = employee_details.add_run(employee_email)
        email_link.font.size = Pt(8)
        email_link.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for email
        
        # Add spacing before signature
        doc.add_paragraph('\n')
        
        # Add signature section at the bottom
        signature_section = doc.add_paragraph()
        signature_section.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add "For COMPANY NAME" text
        company_name = data.get('company', {}).get('name', '').upper()
        for_company = signature_section.add_run(f"For {company_name}")
        for_company.font.size = Pt(11)
        signature_section.add_run('\n\n')  # Add some space
        
        # Add company seal image if available
        company_seal = data.get('company', {}).get('seal_image_url', '')  # Changed from 'seal_image' to 'seal_image_url'
        print("Company Seal Data:", company_seal)  # Debug print
        if company_seal:
            try:
                # Get the full path to the image
                seal_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), company_seal.lstrip('/'))
                
                if os.path.exists(seal_path):
                    # Add the image to the document with specific size
                    signature_section.add_run().add_picture(seal_path, width=Inches(1.1), height=Inches(1.1))  # Slightly larger than 80px
                    signature_section.add_run('\n')  # Add space after seal
                else:
                    print(f"Seal image file not found at path: {seal_path}")
            except Exception as e:
                print(f"Error adding company seal: {str(e)}")
                import traceback
                print(traceback.format_exc())  # Print full error traceback
        else:
            print("No company seal image found in data")  # Debug print
        
        # Add Authorized Signatory text
        signature_section.add_run('\n')  # Add extra space before text
        auth_signatory = signature_section.add_run("Authorized Signatory")
        auth_signatory.font.size = Pt(11)
        
        # Save the document
        filename = f"quote_{quotation_data['ref_number']}.docx"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        doc.save(filepath)
        
        return jsonify({
            'success': True,
            'message': 'Quotation generated successfully',
            'filename': filename
        })
        
    except Exception as e:
        print("Error generating quote:", str(e))  # Debug log
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/hsn/<hsn_code>/gst', methods=['GET'])
def get_gst_percentage(hsn_code):
    try:
        # Comprehensive HSN-GST mapping for chemicals, lab equipment and related products
        gst_mapping = {
            # Chapter 28: Inorganic chemicals
            '28': 18,    # All inorganic chemicals
            '2801': 18,  # Halogens (fluorine, chlorine, bromine and iodine)
            '2802': 18,  # Sulphur, sublimed or precipitated; colloidal sulphur
            '2803': 18,  # Carbon (carbon blacks and other forms of carbon)
            '2804': 18,  # Hydrogen, rare gases and other non-metals
            '2805': 18,  # Alkali or alkaline-earth metals
            '2806': 18,  # Hydrogen chloride and chlorosulphuric acid
            '2807': 18,  # Sulphuric acid and oleum
            '2808': 18,  # Nitric acid; sulphonitric acids
            '2809': 18,  # Phosphorus pentoxide; phosphoric acid
            '2810': 18,  # Oxides of boron; boric acids
            '2811': 18,  # Other inorganic acids
            '2812': 18,  # Halides and halide oxides of non-metals
            '2813': 18,  # Sulphides of non-metals
            '2814': 18,  # Ammonia, anhydrous or in aqueous solution
            '2815': 18,  # Sodium/Potassium hydroxide; peroxides of sodium/potassium
            '2816': 18,  # Hydroxide and peroxide of magnesium
            '2817': 18,  # Zinc oxide; zinc peroxide
            '2818': 18,  # Artificial corundum; aluminum oxide; aluminum hydroxide
            '2819': 18,  # Chromium oxides and hydroxides
            '2820': 18,  # Manganese oxides
            '2821': 18,  # Iron oxides and hydroxides
            '2822': 18,  # Cobalt oxides and hydroxides
            '2823': 18,  # Titanium oxides
            '2824': 18,  # Lead oxides
            '2825': 18,  # Hydrazine, hydroxylamine, inorganic bases
            '2826': 18,  # Fluorides; fluorosilicates, fluoroaluminates
            '2827': 18,  # Chlorides, chloride oxides
            '2828': 18,  # Hypochlorites; commercial calcium hypochlorite
            '2829': 18,  # Chlorates and perchlorates
            '2830': 18,  # Sulphides; polysulphides
            '2831': 18,  # Dithionites and sulphoxylates
            '2832': 18,  # Sulphites; thiosulphates
            '2833': 18,  # Sulphates; alums; peroxosulphates
            '2834': 18,  # Nitrites; nitrates
            '2835': 18,  # Phosphinates, phosphonates, phosphates
            '2836': 18,  # Carbonates; peroxocarbonates
            '2837': 18,  # Cyanides, cyanide oxides
            '2838': 18,  # Fulminates, cyanates and thiocyanates
            '2839': 18,  # Silicates; commercial alkali metal silicates
            '2840': 18,  # Borates; peroxoborates
            '2841': 18,  # Salts of oxometallic/peroxometallic acids
            '2842': 18,  # Other salts of inorganic acids/peroxoacids
            '2843': 18,  # Colloidal precious metals
            '2844': 18,  # Radioactive chemical elements
            '2845': 18,  # Isotopes and their compounds
            '2846': 18,  # Compounds of rare-earth metals
            '2847': 18,  # Hydrogen peroxide
            '2848': 18,  # Phosphides
            '2849': 18,  # Carbides
            '2850': 18,  # Hydrides, nitrides, azides, silicides
            '2851': 18,  # Other inorganic compounds

            # Chapter 29: Organic Chemicals
            '29': 18,    # All organic chemicals
            '2901': 18,  # Acyclic hydrocarbons
            '2902': 18,  # Cyclic hydrocarbons
            '2903': 18,  # Halogenated derivatives of hydrocarbons
            '2904': 18,  # Sulphonated, nitrated derivatives
            '2905': 18,  # Acyclic alcohols and their derivatives
            '2906': 18,  # Cyclic alcohols and their derivatives
            '2907': 18,  # Phenols; phenol-alcohols
            '2908': 18,  # Derivatives of phenols
            '2909': 18,  # Ethers, ether-alcohols, ether-phenols
            '2910': 18,  # Epoxides, epoxyalcohols, epoxyphenols
            '2911': 18,  # Acetals and hemiacetals
            '2912': 18,  # Aldehydes
            '2913': 18,  # Halogenated, sulphonated derivatives of aldehydes
            '2914': 18,  # Ketones and quinones
            '2915': 18,  # Saturated acyclic monocarboxylic acids
            '2916': 18,  # Unsaturated acyclic monocarboxylic acids
            '2917': 18,  # Polycarboxylic acids
            '2918': 18,  # Carboxylic acids with additional oxygen function
            '2919': 18,  # Phosphoric esters and their salts
            '2920': 18,  # Esters of other inorganic acids
            '2921': 18,  # Amine-function compounds
            '2922': 18,  # Oxygen-function amino-compounds
            '2923': 18,  # Quaternary ammonium salts and hydroxides
            '2924': 18,  # Carboxyamide-function compounds
            '2925': 18,  # Carboxyimide-function compounds
            '2926': 18,  # Nitrile-function compounds
            '2927': 18,  # Diazo-, azo- or azoxy-compounds
            '2928': 18,  # Organic derivatives of hydrazine
            '2929': 18,  # Compounds with other nitrogen function
            '2930': 18,  # Organo-sulphur compounds
            '2931': 18,  # Other organo-inorganic compounds
            '2932': 18,  # Heterocyclic compounds with oxygen
            '2933': 18,  # Heterocyclic compounds with nitrogen
            '2934': 18,  # Nucleic acids and their salts
            '2935': 18,  # Sulphonamides
            '2936': 18,  # Provitamins and vitamins
            '2937': 18,  # Hormones
            '2938': 18,  # Glycosides
            '2939': 18,  # Vegetable alkaloids
            '2940': 18,  # Sugars, chemically pure
            '2941': 18,  # Antibiotics
            '2942': 18,  # Other organic compounds

            # Chapter 38: Miscellaneous chemical products
            '38': 18,    # All miscellaneous chemical products
            '3801': 18,  # Artificial graphite; preparations
            '3802': 18,  # Activated carbon; activated natural products
            '3803': 18,  # Tall oil
            '3804': 18,  # Residual lyes from wood pulp
            '3805': 18,  # Gum, wood or sulphate turpentine
            '3806': 18,  # Rosin and resin acids
            '3807': 18,  # Wood tar; wood tar oils
            '3808': 18,  # Insecticides, fungicides
            '3809': 18,  # Finishing agents, dye carriers
            '3810': 18,  # Pickling preparations for metal surfaces
            '3811': 18,  # Anti-knock preparations
            '3812': 18,  # Prepared rubber accelerators
            '3813': 18,  # Preparations for fire-extinguishers
            '3814': 18,  # Organic composite solvents
            '3815': 18,  # Reaction initiators, accelerators
            '3816': 18,  # Refractory cements, mortars
            '3817': 18,  # Mixed alkylbenzenes
            '3818': 18,  # Chemical elements for electronics
            '3819': 18,  # Hydraulic brake fluids
            '3820': 18,  # Anti-freezing preparations
            '3821': 18,  # Prepared culture media
            '3822': 18,  # Diagnostic or laboratory reagents
            '3823': 18,  # Industrial monocarboxylic fatty acids
            '3824': 18,  # Prepared binders; chemical products
            '3825': 18,  # Residual products of chemical industry
            '3826': 18,  # Biodiesel and mixtures

            # Chapter 70: Glass and glassware (Laboratory glassware)
            '7017': 18,  # Laboratory, hygienic or pharmaceutical glassware
            '701710': 18,  # Of fused quartz or other fused silica
            '701720': 18,  # Of other glass having linear coefficient
            '701790': 18,  # Other laboratory glassware

            # Chapter 90: Scientific and laboratory instruments
            '9011': 18,  # Microscopes
            '9012': 18,  # Microscopes other than optical
            '9015': 18,  # Surveying instruments
            '9016': 18,  # Balances of a sensitivity
            '9017': 18,  # Drawing, marking-out instruments
            '9018': 12,  # Medical instruments and appliances
            '9019': 12,  # Mechano-therapy appliances
            '9022': 18,  # X-ray apparatus
            '9023': 18,  # Instruments, apparatus and models
            '9024': 18,  # Machines for testing materials
            '9025': 18,  # Hydrometers, thermometers
            '9026': 18,  # Instruments for measuring flow
            '9027': 18,  # Instruments for physical/chemical analysis
            '9028': 18,  # Gas, liquid or electricity meters
            '9029': 18,  # Revolution counters, taximeters
            '9030': 18,  # Oscilloscopes, spectrum analyzers
            '9031': 18,  # Measuring or checking instruments
            '9032': 18,  # Automatic regulating instruments
            '9033': 18,  # Parts and accessories for machines
        }
        
        # Try to match HSN code at different levels (8, 6, 4, 2 digits)
        hsn_variants = [
            hsn_code,  # Full 8-digit code
            hsn_code[:6] if len(hsn_code) >= 6 else None,  # First 6 digits
            hsn_code[:4] if len(hsn_code) >= 4 else None,  # First 4 digits
            hsn_code[:2] if len(hsn_code) >= 2 else None   # First 2 digits
        ]
        
        # Try each variant, from most specific to least specific
        for variant in hsn_variants:
            if variant and variant in gst_mapping:
                return jsonify({
                    'success': True,
                    'gst_percentage': gst_mapping[variant]
                })
        
        # If no match found
        return jsonify({
            'success': False,
            'error': 'GST percentage not found for this HSN code'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# Item routes
@app.route('/api/items', methods=['GET'])
def get_items():
    try:
        items = supabase.table('items').select('*').execute()
        return jsonify({
            "success": True,
            "data": [Item.from_db(item) for item in items.data]
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/items', methods=['POST'])
def create_item():
    try:
        data = request.json
        item = supabase.table('items').insert(data).execute()
        return jsonify({
            "success": True,
            "data": Item.from_db(item.data[0])
        }), 201
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/items/<int:item_id>', methods=['PUT'])
def update_item(item_id):
    try:
        data = request.json
        item = supabase.table('items').update(data).eq('id', item_id).execute()
        if not item.data:
            return jsonify({"success": False, "error": "Item not found"}), 404
        return jsonify({
            "success": True,
            "data": Item.from_db(item.data[0])
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/items/<int:item_id>', methods=['DELETE'])
def delete_item(item_id):
    try:
        # Check if item is used in any quotations before deleting
        # You might want to add this check when quotations are implemented
        item = supabase.table('items').delete().eq('id', item_id).execute()
        if not item.data:
            return jsonify({"success": False, "error": "Item not found"}), 404
            
        return jsonify({
            "success": True,
            "message": "Item deleted successfully"
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# Error handlers
@app.errorhandler(404)
def not_found_error(error):
    return jsonify({
        "success": False,
        "error": "Resource not found"
    }), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({
        "success": False,
        "error": "Internal server error"
    }), 500

@app.route('/api/generate-quotation', methods=['POST'])
def generate_quotation():
    try:
        data = request.json
        
        # First, save the quotation to the database
        company_id = data.get('company', {}).get('id')
        client_id = data.get('client', {}).get('id')
        employee_id = data.get('employee', {}).get('id')  # Get employee ID
        
        # Generate reference number
        company_response = supabase.table('companies').select('*').eq('id', company_id).execute()
        if not company_response.data:
            return jsonify({"success": False, "error": "Company not found"}), 404
            
        company = company_response.data[0]
        new_number = company.get('last_quote_number', 0) + 1
        ref_number = company.get('ref_format', 'QUOTE-{YYYY}-{NUM}').format(
            YYYY=datetime.now().year,
            NUM=str(new_number).zfill(4)
        )
        
        # Prepare quotation data
        quotation_data = {
            'company_id': company_id,
            'client_id': client_id,
            'ref_number': ref_number,
            'date': datetime.utcnow().isoformat(),
            'items': data.get('items', []),
            'total': data.get('grandTotal', 0)
        }
        
        # Create quotation in database
        quotation_response = supabase.table('quotations').insert(quotation_data).execute()
        
        if not quotation_response.data:
            raise Exception("Failed to create quotation")
            
        # Update company's last quote number
        company_update = supabase.table('companies').update({
            'last_quote_number': new_number
        }).eq('id', company_id).execute()
        
        if not company_update.data:
            print("Warning: Failed to update company's last quote number")
        
        # Now proceed with document generation
        doc = Document()
        
        # Set very narrow margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.3)    # 0.3 inch top margin
            section.bottom_margin = Inches(0.3)  # 0.3 inch bottom margin
            section.left_margin = Inches(0.3)    # 0.3 inch left margin
            section.right_margin = Inches(0.3)   # 0.3 inch right margin
        
        # Set default font size for the document
        style = doc.styles['Normal']
        style.font.size = Pt(8)  # Changed from 9 to 8
        
        # Add header table
        header_table = doc.add_table(rows=4, cols=1)
        header_table.style = 'Table Grid'
        
        # Company Name Cell
        company_cell = header_table.rows[0].cells[0]
        company_name = company_cell.paragraphs[0]
        company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_name.add_run(data.get('company', {}).get('name', '').upper())
        company_name.runs[0].font.size = Pt(16)
        company_name.runs[0].font.bold = True
        company_name.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Address Cell
        address_cell = header_table.rows[1].cells[0]
        address = address_cell.paragraphs[0]
        address.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address.add_run(data.get('company', {}).get('address', '').upper())
        address.runs[0].font.size = Pt(11)
        address.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Email and Phone Cell
        email_cell = header_table.rows[2].cells[0]
        email = email_cell.paragraphs[0]
        email.alignment = WD_ALIGN_PARAGRAPH.CENTER
        email.add_run(f"Email:- {data.get('company', {}).get('email', '')} {data.get('company', {}).get('phone', '')}")
        email.runs[0].font.size = Pt(11)
        email.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Tax Info Cell
        tax_cell = header_table.rows[3].cells[0]
        tax_info = tax_cell.paragraphs[0]
        tax_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pan = data.get('company', {}).get('pan_number', '')  # Changed from 'pan' to 'pan_number'
        gst = data.get('company', {}).get('gst_number', '')  # Changed from 'gst' to 'gst_number'
        tax_text = f"PAN NO.: {pan} | GST NO.: {gst}"
        tax_info.add_run(tax_text)
        tax_info.runs[0].font.size = Pt(11)
        tax_info.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Set background color for all cells and remove borders
        for row in header_table.rows:
            for cell in row.cells:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4F8C"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
                # Remove cell borders
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>' +
                                    '<w:top w:val="nil"/>' +
                                    '<w:left w:val="nil"/>' +
                                    '<w:bottom w:val="nil"/>' +
                                    '<w:right w:val="nil"/>' +
                                    '</w:tcBorders>')
                tcPr.append(tcBorders)

        # Add QUOTATION/PERFORMA INVOICE title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run('\nQUOTATION/PERFORMA INVOICE\n')
        title.runs[0].font.bold = True
        title.runs[0].font.size = Pt(12)

        # Add reference number and date
        ref_date = doc.add_table(rows=1, cols=2)
        ref_date.autofit = True
        ref_cell = ref_date.cell(0, 0)
        ref_cell.text = f"Ref No: {data.get('refNumber', '')}"
        date_cell = ref_date.cell(0, 1)
        date_cell.text = f"Date: {data.get('quotationDate', '')}"
        date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add spacing after header
        doc.add_paragraph()

        # Set default font size for the document
        style = doc.styles['Normal']
        style.font.size = Pt(9)

        # Add client details
        to_table = doc.add_table(rows=2, cols=1)
        to_table.style = 'Table Grid'
        
        # To Cell with blue background
        to_cell = to_table.rows[0].cells[0]
        to_paragraph = to_cell.paragraphs[0]
        to_paragraph.add_run('To')
        to_paragraph.runs[0].font.bold = True
        to_paragraph.runs[0].font.size = Pt(9)
        to_paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Set blue background for To cell
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4F8C"/>')
        to_cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Client details cell
        client = data.get('client', {})
        details_cell = to_table.rows[1].cells[0]
        details_paragraph = details_cell.paragraphs[0]
        details_paragraph.add_run(f"{client.get('business_name', '')}\n")
        details_paragraph.add_run(f"{client.get('address', '')}\n")
        details_paragraph.add_run(f"Kind Attn: {client.get('name', '')} | Tel: {client.get('phone', '')} | Email: {client.get('email', '')}")
        
        # Remove borders from both cells
        for row in to_table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>' +
                                    '<w:top w:val="nil"/>' +
                                    '<w:left w:val="nil"/>' +
                                    '<w:bottom w:val="nil"/>' +
                                    '<w:right w:val="nil"/>' +
                                    '</w:tcBorders>')
                tcPr.append(tcBorders)
        
        # Add spacing after client details
        doc.add_paragraph()
        
        # Add greeting text
        greeting = doc.add_paragraph()
        greeting.add_run("Dear Sir/Madam,\n")
        greeting.add_run("Thank you for your enquiry. We are pleased to quote our best prices as under:\n")
        
        # Add spacing after greeting
        doc.add_paragraph()
        
        # Add items table
        doc.add_paragraph().add_run('Items:').bold = True
        table = doc.add_table(rows=1, cols=14)
        table.style = 'Table Grid'
        table.allow_autofit = True  # Enable auto-fitting
        
        # Set header row with smaller font
        header_cells = table.rows[0].cells
        headers = ['S.No', 'Cat No.', 'Description', 'Pack Size', 'HSN Code', 'Qty', 'Unit Rate', 'Discounted Price', 'Expanded Price', 'GST %', 'GST', 'Total Value', 'Lead Time', 'Brand']
        for i, text in enumerate(headers):
            header_cells[i].text = text
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)  # Set header font size
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add item rows
        items = data.get('items', [])
        for idx, item in enumerate(items, 1):
            row = table.add_row().cells
            row[0].text = str(idx)  # S.No
            row[1].text = item.get('catalogue_id', '')  # Cat No.
            row[2].text = item.get('description', '')  # Description
            row[3].text = item.get('pack_size', '')  # Pack Size
            row[4].text = item.get('hsn', '')  # HSN Code
            row[5].text = str(item.get('quantity', ''))  # Qty
            row[6].text = f"₹{item.get('unit_rate', 0):.2f}"  # Unit Rate
            
            # Calculate discounted price
            unit_rate = float(item.get('unit_rate', 0))
            discount = float(item.get('discount_percentage', 0))
            discounted_price = unit_rate * (1 - discount/100)
            row[7].text = f"₹{discounted_price:.2f}"  # Discounted Price
            
            # Calculate expanded price (discounted price * quantity)
            quantity = float(item.get('quantity', 0))
            expanded_price = discounted_price * quantity
            row[8].text = f"₹{expanded_price:.2f}"  # Expanded Price
            
            row[9].text = f"{item.get('gst_percentage', 0)}%"  # GST %
            row[10].text = f"₹{item.get('gst_value', 0):.2f}"  # GST
            row[11].text = f"₹{item.get('total', 0):.2f}"  # Total Value
            row[12].text = item.get('lead_time', '')  # Lead Time
            row[13].text = item.get('brand', '')  # Changed from 'make' to 'brand'
            
            # Center align all cells
            for cell in row:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set optimized column widths for better fit
        for cell in table.columns[0].cells:  # S.No
            cell.width = Inches(0.2)
        for cell in table.columns[1].cells:  # Cat No.
            cell.width = Inches(0.6)
        for cell in table.columns[2].cells:  # Description
            cell.width = Inches(1.8)
        for cell in table.columns[3].cells:  # Pack Size
            cell.width = Inches(0.4)
        for cell in table.columns[4].cells:  # HSN Code
            cell.width = Inches(0.6)
        for cell in table.columns[5].cells:  # Qty
            cell.width = Inches(0.3)
        for cell in table.columns[6].cells:  # Unit Rate
            cell.width = Inches(0.6)
        for cell in table.columns[7].cells:  # Discounted Price
            cell.width = Inches(0.6)
        for cell in table.columns[8].cells:  # Expanded Price
            cell.width = Inches(0.6)
        for cell in table.columns[9].cells:  # GST %
            cell.width = Inches(0.4)
        for cell in table.columns[10].cells:  # GST
            cell.width = Inches(0.6)
        for cell in table.columns[11].cells:  # Total Value
            cell.width = Inches(0.6)
        for cell in table.columns[12].cells:  # Lead Time
            cell.width = Inches(0.6)
        for cell in table.columns[13].cells:  # Brand
            cell.width = Inches(0.5)
        
        # Add totals
        doc.add_paragraph()
        totals = doc.add_paragraph()
        totals.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        totals.add_run(f"Sub Total: ₹{data.get('subTotal', 0):.2f}\n")
        totals.add_run(f"Total GST: ₹{data.get('totalGST', 0):.2f}\n")
        grand_total = totals.add_run(f"Grand Total: ₹{data.get('grandTotal', 0):.2f}")
        grand_total.bold = True
        
        # Add terms and conditions
        terms_table = doc.add_table(rows=1, cols=1)
        terms_table.style = 'Table Grid'
        
        # Terms & Conditions header cell with blue background
        header_cell = terms_table.rows[0].cells[0]
        header_paragraph = header_cell.paragraphs[0]
        header_run = header_paragraph.add_run('Terms & Conditions')
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(255, 255, 255)
        header_run.font.size = Pt(11)
        
        # Set blue background for header cell
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4F8C"/>')
        header_cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Remove borders from the header cell
        tcPr = header_cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>' +
                            '<w:top w:val="nil"/>' +
                            '<w:left w:val="nil"/>' +
                            '<w:bottom w:val="nil"/>' +
                            '<w:right w:val="nil"/>' +
                            '</w:tcBorders>')
        tcPr.append(tcBorders)
        
        # Add payment terms and other terms as bullet points
        terms_list = doc.add_paragraph()
        terms_list.style = doc.styles['Normal']
        terms_list.paragraph_format.space_before = Pt(6)
        
        # Add payment terms
        payment_term = terms_list.add_run(f"1) {data.get('paymentTerms', '')}\n")
        payment_term.font.size = Pt(8)
        
        # Add fixed terms with numbering
        for idx, term in enumerate(data.get('fixedTerms', []), 2):
            term_text = terms_list.add_run(f"{idx}) {term}\n")
            term_text.font.size = Pt(8)
        
        # Add spacing
        doc.add_paragraph('\n')
        
        # Add Bank Details section
        bank_table = doc.add_table(rows=1, cols=1)
        bank_table.style = 'Table Grid'
        
        # Bank Details header cell with blue background
        bank_header_cell = bank_table.rows[0].cells[0]
        bank_header_paragraph = bank_header_cell.paragraphs[0]
        bank_header_run = bank_header_paragraph.add_run('Bank Details')
        bank_header_run.font.bold = True
        bank_header_run.font.color.rgb = RGBColor(255, 255, 255)
        bank_header_run.font.size = Pt(11)
        
        # Set blue background for header cell
        bank_shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4F8C"/>')
        bank_header_cell._tc.get_or_add_tcPr().append(bank_shading_elm)
        
        # Remove borders from the header cell
        bank_tcPr = bank_header_cell._tc.get_or_add_tcPr()
        bank_tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>' +
                            '<w:top w:val="nil"/>' +
                            '<w:left w:val="nil"/>' +
                            '<w:bottom w:val="nil"/>' +
                            '<w:right w:val="nil"/>' +
                            '</w:tcBorders>')
        bank_tcPr.append(bank_tcBorders)
        
        # Add bank details content
        bank_details = doc.add_paragraph()
        bank_details.style = doc.styles['Normal']
        bank_details.paragraph_format.space_before = Pt(6)
        bank_details_text = bank_details.add_run(f"HDFC BANK LTD. Account No: {data.get('company', {}).get('account_number', '')} ; NEFT/RTGS IFCS : {data.get('company', {}).get('ifsc_code', '')} Branch code:{data.get('company', {}).get('branch_code', '')} ; Micro code : {data.get('company', {}).get('micro_code', '')} ;Account type: Current account")
        bank_details_text.font.size = Pt(8)
        
        # Add spacing before quotation created by section
        doc.add_paragraph('\n')
        
        # Add Quotation Created By section
        created_by_table = doc.add_table(rows=1, cols=1)
        created_by_table.style = 'Table Grid'
        
        # Created By header cell with blue background
        created_by_header_cell = created_by_table.rows[0].cells[0]
        created_by_header_paragraph = created_by_header_cell.paragraphs[0]
        created_by_header_run = created_by_header_paragraph.add_run('Quotation Created By')
        created_by_header_run.font.bold = True
        created_by_header_run.font.color.rgb = RGBColor(255, 255, 255)
        created_by_header_run.font.size = Pt(11)
        
        # Set blue background for header cell
        created_by_shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4F8C"/>')
        created_by_header_cell._tc.get_or_add_tcPr().append(created_by_shading_elm)
        
        # Remove borders from the header cell
        created_by_tcPr = created_by_header_cell._tc.get_or_add_tcPr()
        created_by_tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>' +
                            '<w:top w:val="nil"/>' +
                            '<w:left w:val="nil"/>' +
                            '<w:bottom w:val="nil"/>' +
                            '<w:right w:val="nil"/>' +
                            '</w:tcBorders>')
        created_by_tcPr.append(created_by_tcBorders)
        
        # Add employee details content
        employee_details = doc.add_paragraph()
        employee_details.style = doc.styles['Normal']
        employee_details.paragraph_format.space_before = Pt(6)
        employee_name = data.get('employee', {}).get('name', '')
        employee_phone = data.get('employee', {}).get('phone_number', '')  # Changed from mobile to phone_number
        employee_email = data.get('employee', {}).get('email', '')
        
        employee_details.add_run(f"{employee_name}\n").font.size = Pt(8)
        employee_details.add_run(f"Mobile: {employee_phone}\n").font.size = Pt(8)  # Using the new employee_phone variable
        email_run = employee_details.add_run(f"Email: ")
        email_run.font.size = Pt(8)
        email_link = employee_details.add_run(employee_email)
        email_link.font.size = Pt(8)
        email_link.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for email
        
        # Add spacing before signature
        doc.add_paragraph('\n')
        
        # Add signature section at the bottom
        signature_section = doc.add_paragraph()
        signature_section.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add "For COMPANY NAME" text
        company_name = data.get('company', {}).get('name', '').upper()
        for_company = signature_section.add_run(f"For {company_name}")
        for_company.font.size = Pt(11)
        signature_section.add_run('\n\n')  # Add some space
        
        # Add company seal image if available
        company_seal = data.get('company', {}).get('seal_image_url', '')  # Changed from 'seal_image' to 'seal_image_url'
        print("Company Seal Data:", company_seal)  # Debug print
        if company_seal:
            try:
                # Get the full path to the image
                seal_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), company_seal.lstrip('/'))
                
                if os.path.exists(seal_path):
                    # Add the image to the document with specific size
                    signature_section.add_run().add_picture(seal_path, width=Inches(1.1), height=Inches(1.1))  # Slightly larger than 80px
                    signature_section.add_run('\n')  # Add space after seal
                else:
                    print(f"Seal image file not found at path: {seal_path}")
            except Exception as e:
                print(f"Error adding company seal: {str(e)}")
                import traceback
                print(traceback.format_exc())  # Print full error traceback
        else:
            print("No company seal image found in data")  # Debug print
        
        # Add Authorized Signatory text
        signature_section.add_run('\n')  # Add extra space before text
        auth_signatory = signature_section.add_run("Authorized Signatory")
        auth_signatory.font.size = Pt(11)
        
        # Save the document
        filename = f"quotation_{data.get('refNumber', 'temp').replace('/', '_')}.docx"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        doc.save(filepath)
        
        return jsonify({
            'success': True,
            'message': 'Quotation generated successfully',
            'filename': filename
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/download-quotation/<filename>', methods=['GET'])
def download_quotation(filename):
    try:
        return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)
    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        }), 404

# Add debug route to check environment variables
@app.route('/api/debug/config', methods=['GET'])
def debug_config():
    return jsonify({
        "supabase_url_exists": bool(os.getenv('SUPABASE_URL')),
        "supabase_key_exists": bool(os.getenv('SUPABASE_KEY')),
        "secret_key_exists": bool(os.getenv('SECRET_KEY'))
    }), 200

@app.route('/api/quotations/<int:quotation_id>', methods=['GET'])
def get_quotation(quotation_id):
    try:
        print(f"Fetching quotation with ID: {quotation_id}")  # Debug log
        
        # First, check if quotation exists
        quotation = supabase.table('quotations') \
            .select('*') \
            .eq('id', quotation_id) \
            .execute()
            
        if not quotation.data:
            return jsonify({
                "success": False,
                "error": "Quotation not found"
            }), 404

        quotation_data = quotation.data[0]
        print(f"Found quotation: {quotation_data}")  # Debug log

        # Fetch company details
        company = supabase.table('companies') \
            .select('name, gst_number') \
            .eq('id', quotation_data['company_id']) \
            .execute()
            
        print(f"Company data: {company.data}")  # Debug log

        # Fetch client details
        client = supabase.table('clients') \
            .select('*') \
            .eq('id', quotation_data['client_id']) \
            .execute()
            
        print(f"Client data: {client.data}")  # Debug log

        # Process the data
        processed_quotation = {
            'id': quotation_data['id'],
            'ref_number': quotation_data['ref_number'],
            'date': quotation_data['date'],
            'company': company.data[0]['name'] if company.data else None,
            'company_gst': company.data[0]['gst_number'] if company.data else None,
            'client': client.data[0]['name'] if client.data else None,
            'client_contact': client.data[0]['contact'] if client.data and 'contact' in client.data[0] else None,  # Using 'contact' as a fallback
            'total_amount': float(quotation_data['total']) if quotation_data.get('total') else 0,
            'items': []
        }

        # Process items from the quotation data itself since they're stored in JSON
        if quotation_data.get('items'):
            for item in quotation_data['items']:
                processed_item = {
                    'id': item['id'],
                    'description': item['description'],
                    'quantity': item['quantity'],
                    'price': float(item['unit_rate']) if item.get('unit_rate') else 0,
                    'gst_percentage': item['gst_percentage'],
                    'total': float(item['total']) if item.get('total') else 0
                }
                processed_quotation['items'].append(processed_item)

        print(f"Processed quotation: {processed_quotation}")  # Debug log

        return jsonify({
            "success": True,
            "data": processed_quotation
        })

    except Exception as e:
        import traceback
        print(f"Error in get_quotation: {str(e)}")
        print(f"Traceback: {traceback.format_exc()}")  # Print full error traceback
        return jsonify({
            "success": False,
            "error": f"Failed to fetch quotation details: {str(e)}"
        }), 500

@app.route('/api/quotations/<int:quotation_id>/download', methods=['GET'])
def download_quotation_by_id(quotation_id):
    try:
        # First, get the quotation details
        quotation = supabase.table('quotations') \
            .select('*') \
            .eq('id', quotation_id) \
            .execute()
            
        if not quotation.data:
            return jsonify({
                "success": False,
                "error": "Quotation not found"
            }), 404

        quotation_data = quotation.data[0]
        
        # Generate the Word document
        doc = DocxTemplate("templates/quote_template.docx")
        
        # Get company and client details
        company = supabase.table('companies') \
            .select('*') \
            .eq('id', quotation_data['company_id']) \
            .execute()
            
        client = supabase.table('clients') \
            .select('*') \
            .eq('id', quotation_data['client_id']) \
            .execute()

        # Prepare context for the template
        context = {
            'quotation_number': quotation_data['ref_number'],
            'date': quotation_data.get('date', '').split('T')[0],
            'company_name': company.data[0]['name'] if company.data else '',
            'company_gst': company.data[0]['gst_number'] if company.data else '',
            'client_name': client.data[0]['name'] if client.data else '',
            'items': quotation_data.get('items', []),
            'total_amount': quotation_data.get('total', 0)
        }

        # Render the template
        doc.render(context)
        
        # Save the document
        filename = f"quotation_{quotation_data['ref_number']}.docx"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        doc.save(filepath)

        return send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        import traceback
        print(f"Error in download_quotation: {str(e)}")
        print(f"Traceback: {traceback.format_exc()}")
        return jsonify({
            "success": False,
            "error": f"Failed to generate quotation: {str(e)}"
        }), 500

if __name__ == '__main__':
    print("Starting Flask server on port 5000...")
    app.run(host='0.0.0.0', port=5000, debug=True) 