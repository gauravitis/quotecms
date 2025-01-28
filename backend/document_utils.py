from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import requests
from io import BytesIO

def add_image_from_url(doc, url, width=2.0):
    """Add an image from URL to the document"""
    try:
        response = requests.get(url)
        image_stream = BytesIO(response.content)
        doc.add_picture(image_stream, width=Inches(width))
    except Exception as e:
        print(f"Error adding image: {str(e)}")
        # Add a placeholder text instead
        doc.add_paragraph("[Image not available]")

def create_quotation_template():
    """Create a new quotation template with placeholders"""
    doc = Document()
    
    # Add header with company information
    header = doc.add_heading('QUOTATION', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add reference number
    doc.add_paragraph('REFERENCE NUMBER: {{ ref_number }}')
    doc.add_paragraph('DATE: {{ date }}')
    
    # Add spacing
    doc.add_paragraph()
    
    # Client Information section
    doc.add_heading('Client Information', level=1)
    doc.add_paragraph('CLIENT NAME: {{ client_name }}')
    doc.add_paragraph('EMAIL: {{ client_email }}')
    
    # Add spacing
    doc.add_paragraph()
    
    # Items section
    doc.add_heading('Items', level=1)
    doc.add_paragraph('{% for item in items %}')
    p = doc.add_paragraph()
    p.add_run('- {{ item.name }} ').bold = True
    p.add_run('(Qty: {{ item.qty }}) - Price: ${{ item.price }}')
    doc.add_paragraph('{% endfor %}')
    
    # Add spacing
    doc.add_paragraph()
    
    # Total section
    doc.add_heading('Total', level=1)
    doc.add_paragraph('SUBTOTAL: ${{ subtotal }}')
    doc.add_paragraph('TAX ({{ tax_rate }}%): ${{ tax_amount }}')
    doc.add_paragraph('TOTAL: ${{ total }}')
    
    # Add spacing
    doc.add_paragraph()
    
    # Company seal and signature
    doc.add_paragraph('Company Seal:')
    doc.add_paragraph('[SEAL_IMAGE_PLACEHOLDER]')  # Will be replaced with actual image
    
    # Save the template
    template_path = os.path.join('templates', 'quote_template.docx')
    doc.save(template_path)
    return template_path

def generate_quotation_doc(template_path, data):
    """Generate a quotation document from template and data"""
    doc = Document(template_path)
    
    # Replace text placeholders
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key != 'seal_image_url':  # Handle image separately
                placeholder = '{{ ' + key + ' }}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Handle image placeholder
        if '[SEAL_IMAGE_PLACEHOLDER]' in paragraph.text:
            paragraph.text = ''  # Clear the placeholder
            if data.get('seal_image_url'):
                add_image_from_url(doc, data['seal_image_url'])
    
    return doc

if __name__ == '__main__':
    # Create the template when this script is run directly
    create_quotation_template() 